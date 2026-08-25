"""tasks/remediation-plan.md R3 — DocxParser (document-processing.md §3.2)."""

import io

import pytest

from app.document_processing.base import CorruptFileError
from app.document_processing.docx_parser import DocxParser


def _build_docx(
    paragraphs: list[str], table_rows: list[list[str]] | None = None
) -> bytes:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


@pytest.fixture
def parser() -> DocxParser:
    return DocxParser()


def test_sniff_matches_zip_container_header(parser):
    assert parser.sniff_matches(b"PK\x03\x04 rest of header") is True


def test_parse_extracts_paragraph_text_in_order(parser):
    data = _build_docx(["First paragraph.", "Second paragraph."])
    result = parser.parse(data)

    assert "First paragraph." in result.full_text
    assert "Second paragraph." in result.full_text
    assert result.full_text.index("First paragraph.") < result.full_text.index(
        "Second paragraph."
    )
    # DOCX has no file-format-level pagination (document-processing.md §3.2).
    assert result.page_breaks is None
    assert result.page_count is None


def test_parse_flattens_table_rows_as_joined_cell_text(parser):
    data = _build_docx(
        ["Intro paragraph."],
        table_rows=[["Name", "Score"], ["Alice", "95"]],
    )
    result = parser.parse(data)

    assert "Name | Score" in result.full_text
    assert "Alice | 95" in result.full_text


def test_parse_empty_document_yields_empty_text(parser):
    """No DOCX-specific error for this case — the generic degenerate-input
    path (EmptyDocumentError, rag.md §2) handles it downstream."""
    data = _build_docx([])
    result = parser.parse(data)
    assert result.full_text == ""


def test_parse_corrupt_file_raises_typed_error(parser):
    garbage = b"PK\x03\x04 not a real docx zip structure" + b"X" * 100
    with pytest.raises(CorruptFileError) as exc_info:
        parser.parse(garbage)
    assert exc_info.value.retryable is False
