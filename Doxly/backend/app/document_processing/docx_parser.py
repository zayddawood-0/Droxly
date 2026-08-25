"""
tasks/remediation-plan.md R3 — DOCX parser (document-processing.md §3.2,
decisions.md ADR-014: python-docx).
"""

import io
import zipfile

from app.document_processing.base import CorruptFileError, DocumentParser, ParsedText


class DocxParser(DocumentParser):
    mime_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    def sniff_matches(self, header_bytes: bytes) -> bool:
        # DOCX is a ZIP/OPC container (document-processing.md §1).
        return header_bytes.startswith(b"PK\x03\x04")

    def parse(self, data: bytes) -> ParsedText:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = docx.Document(io.BytesIO(data))
        except (
            PackageNotFoundError,
            zipfile.BadZipFile,
            ValueError,
            KeyError,
            OSError,
        ) as exc:
            raise CorruptFileError() from exc

        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # "Embedded tables are extracted as flattened text (row cells
        # joined) inline with surrounding paragraph order" (§3.2) — this
        # codebase's python-docx usage walks `document.tables` after all
        # paragraphs rather than interleaving true document order, a
        # documented, best-effort simplification (§3.2 already allows
        # "not pixel/layout-perfect" for tables).
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        full_text = "\n\n".join(parts)
        # No page numbers at the file-format level (§3.2) — page_count/
        # page_breaks stay None, matching document_chunks.page_number's
        # nullable column. An empty result (no extractable text at all)
        # is left to the generic degenerate-input path (EmptyDocumentError,
        # rag.md §2) rather than a DOCX-specific error the spec doesn't define.
        return ParsedText(full_text=full_text, page_breaks=None, page_count=None)
